# -*- coding: utf-8 -*-
"""生成《给师叔的 GitHub Issue 建议（中文版）》Word 文档"""
import sys
from pathlib import Path
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = Path(r"D:\股票分析项目\2.0版\项目规划\GitHub_Issue_给师叔的建议_中文版.docx")

doc = Document()

# 全局默认字体：微软雅黑（中西文）
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, 16, True, (0x1E, 0x3A, 0x8A))
    p.space_after = Pt(6)

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, 13, True, (0x25, 0x63, 0xEB))
    p.space_before = Pt(10)

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, 11.5, True, (0x2D, 0x37, 0x48))
    p.space_before = Pt(6)

def body(text, indent=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, 11)
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    r0 = p.add_run("• ")
    set_font(r0, 11, True)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        set_font(rb, 11, True)
    r = p.add_run(text)
    set_font(r, 11)

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, 9.5)
    r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    p.paragraph_format.left_indent = Pt(24)

# ================= 文档内容 =================
p = doc.add_paragraph()
r = p.add_run("GitHub Issue 建议（中文版）")
set_font(r, 18, True, (0x1E, 0x3A, 0x8A))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
r = p.add_run("仓库：serenity-chokepoint-investing-enhanced　|　作者：吴宇轩（华南师范大学阿伯丁学院 信息管理与信息系统 二年级）")
set_font(r, 10, False, (0x71, 0x80, 0x96))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

body("师叔您好！老师让我认真研读这个框架，我花了两天逐字读完了全部 6 篇 ADR、README 和 SKILL.md，并做了两只标的的封箱回测（立新能源 001258、美光科技 MU），有一些发现想分享给您。")

h1("A 部分：文档不一致（疑似 v2.0 → v2.1 迭代遗留）")

h2("A1：IR 淘汰阈值 0.5 与 0.3 不一致")
bullet("ADR-0001 的硬约束表写的是 IR < 0.5 拒绝；README / SKILL.md（v2.1）已改为 IR < 0.3 拒绝（0.3~0.5 为弱 alpha、仅小仓位）。", "位置：")
bullet("建议在 ADR-0001 增加一段 “v2.1 更新说明”，指向现行 0.3 阈值，避免后来者被两个数字误导。", "建议：")

h2("A2：仓位调整公式：乘法与加权不一致")
bullet("ADR-0006 用的是乘法公式（Base × (1+调整)×…）；v2.1 正文已改为绝对百分点加减，并设硬地板 max(Base×25%, 0.5%)。", "位置：")
bullet("建议在 ADR-0006 补充 v2.1 更新说明，注明乘法公式保留作历史参考。", "建议：")

h1("B 部分：两个基于回测证据的架构改进建议")

h2("B1：赌注类型应数据化分类，而非仅靠 LLM 判断")
body("ADR-0006 要求 agent 将标的分为 Super Beta / Catalyst Alpha / Event-Driven，但分类依据是 LLM 定性判断。我用同一套信号系统做了对比回测：")
code("立新能源 001258（高波动“妖股”）：持有 +85.3%，信号择时 +163.0%（择时大幅胜出）")
code("美光科技 MU（强趋势牛股）：持有 +423.3%，信号择时 +161.0%（择时大幅跑输）")
body("同一套策略，在两只标的上结果完全相反。这说明最优持仓周期与仓位逻辑取决于标的的统计特征（波动率、动量半衰期、收益自相关）。建议增加一个轻量统计分类器（如 20 日波动率 + 动量半衰期 + ATR），用数据判定赌注类型，与 LLM 定性判断互为校验。")

h2("B2：缺少“趋势环境过滤层”（实测最大的盲区）")
body("美光科技回测中，集成投票的 1 日方向命中率按市场环境分段：")
code("上升段：53.7%　|　下降段：45.2%（低于随机！）　|　震荡段：56.8%")
body("系统在下跌段系统性误判为看涨。建议增加一道“趋势门”：当标的 20/60 日趋势向下时，抑制或降权看多信号。这是我实测中影响最大的一个改进点。")

h1("结尾")
body("以上所有回测产物（报告、JSON、脚本）均可提供。如果方便，我也可以把 A1/A2 的文档修复直接做成 Pull Request。感谢师叔的框架带给我的启发！")
body("此致敬礼", indent=True)
body("吴宇轩", indent=True)

doc.save(OUT)
print("[OK] 中文版 Word 已生成:", OUT)
