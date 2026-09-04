# -*- coding: utf-8 -*-
"""scratch/build_perfect_attachment3.py
High-end typography and clean 2-page layout for Attachment 3 application form.
Conforms to official Chinese administrative and competition typesetting standards:
- Headings & Labels: 黑体 (SimHei), Bold
- Chinese Body: 仿宋 (FangSong) or 宋体 (SimSun)
- Numbers & Latin: Times New Roman
- Checkbox Symbols: Segoe UI Symbol / 宋体
- Circled numbers: 宋体 / Segoe UI Symbol
- Note: 楷体 (KaiTi)
- Strict 2-Page Pagination: Page 1 for Team Info, Page 2 for Strategy & Evidence.
"""

import os, sys, shutil
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

REPO_ROOT = Path(r"d:\R-FinGPTv2（国创版本）")

def format_run(run, font_cn="仿宋", font_en="Times New Roman", font_size=9.0, bold=False, color_rgb=(15, 23, 42)):
    """Apply strict bilingual fonts with eastAsia and Western font separation."""
    run.font.name = font_en
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color_rgb)
    
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:cs'), font_en)

def set_cell_borders_and_padding(cell, top=50, bottom=50, left=80, right=80):
    """Set cell padding in twentieths of a point (dxa). 20 dxa = 1 pt."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="8" w:space="0" w:color="334155"/>'
        f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="334155"/>'
        f'<w:left w:val="single" w:sz="8" w:space="0" w:color="334155"/>'
        f'<w:right w:val="single" w:sz="8" w:space="0" w:color="334155"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_styled_doc(font_cn="仿宋", font_en="Times New Roman", font_title="黑体"):
    doc = docx.Document()
    
    # Configure Normal Style defaults to eliminate Cambria/Calibri fallback
    style = doc.styles['Normal']
    style.font.name = font_en
    style.font.size = Pt(9.0)
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:cs'), font_en)

    for section in doc.sections:
        section.top_margin = Inches(0.50)
        section.bottom_margin = Inches(0.50)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    # Helper to populate cells
    def populate_cell(cell, text, bold=False, font_size=9.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_color=None, is_label=False, line_spacing=1.12, pad_top=50, pad_bottom=50):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(1.2)
        p.paragraph_format.space_after = Pt(1.2)
        p.paragraph_format.line_spacing = line_spacing
        
        c_font = font_title if is_label else font_cn
        lines = text.split("\n")
        for idx, line in enumerate(lines):
            if idx > 0:
                p = cell.add_paragraph()
                p.alignment = align
                p.paragraph_format.space_before = Pt(1.2)
                p.paragraph_format.space_after = Pt(1.2)
                p.paragraph_format.line_spacing = line_spacing
            
            # Special check for checkbox and circled numbers
            special_chars = set("☑□①②③④⑤⑥⑦⑧⑨⑩")
            if any(ch in special_chars for ch in line):
                parts = []
                curr = ""
                curr_is_spec = False
                for ch in line:
                    is_spec = ch in special_chars
                    if not parts and not curr:
                        curr = ch
                        curr_is_spec = is_spec
                    elif is_spec == curr_is_spec:
                        curr += ch
                    else:
                        parts.append((curr, curr_is_spec))
                        curr = ch
                        curr_is_spec = is_spec
                if curr:
                    parts.append((curr, curr_is_spec))
                
                for text_chunk, is_spec in parts:
                    r = p.add_run(text_chunk)
                    if is_spec:
                        format_run(r, font_cn="宋体", font_en="Segoe UI Symbol", font_size=font_size, bold=bold)
                    else:
                        format_run(r, font_cn=c_font, font_en=font_en, font_size=font_size, bold=bold)
            else:
                r = p.add_run(line)
                format_run(r, font_cn=c_font, font_en=font_en, font_size=font_size, bold=bold)

        if bg_color:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            cell._tc.get_or_add_tcPr().append(shading)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders_and_padding(cell, top=pad_top, bottom=pad_bottom, left=70, right=70)

    # -------------------------------------------------------------
    # PAGE 1: Title, Note, and Table 1 (Team & Basic Information)
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    title_p.paragraph_format.line_spacing = 1.15

    r1 = title_p.add_run("中国国际大学生创新大赛（2026）\n")
    format_run(r1, font_cn=font_title, font_en=font_title, font_size=15.0, bold=True, color_rgb=(15, 23, 42))

    r2 = title_p.add_run("产业命题赛道（企业命题组）参赛作品申报表")
    format_run(r2, font_cn=font_title, font_en=font_title, font_size=14.0, bold=True, color_rgb=(15, 23, 42))

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_note.paragraph_format.space_before = Pt(1)
    p_note.paragraph_format.space_after = Pt(2.5)
    p_note.paragraph_format.line_spacing = 1.10
    rn = p_note.add_run("注意！本表项目成员信息次序（自上而下）即为项目成员在项目中的重要程度排序，如项目在校赛阶段获奖，则制作证书时按本表次序进行信息录入，原则上不作修改。")
    format_run(rn, font_cn="楷体", font_en="Times New Roman", font_size=8.5, bold=False, color_rgb=(71, 85, 105))

    # Table 1: 15 rows, 6 columns
    table1 = doc.add_table(rows=15, cols=6)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = False
    set_table_borders(table1)

    for row in table1.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

    # Row 0: 命题企业
    populate_cell(table1.cell(0, 0), "命题企业", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F1F5F9", is_label=True)
    table1.cell(0, 1).merge(table1.cell(0, 5))
    populate_cell(table1.cell(0, 1), "达观数据有限公司 (Datagrand Inc.)", font_size=9.0)

    # Row 1: 命题名称
    populate_cell(table1.cell(1, 0), "命题名称", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F1F5F9", is_label=True)
    table1.cell(1, 1).merge(table1.cell(1, 5))
    populate_cell(table1.cell(1, 1), "面向金融量化投研工作全流程的智能体系统 (命题编号: 新工科/01)", font_size=9.0)

    # Row 2: 参赛对策（项目）名称
    populate_cell(table1.cell(2, 0), "参赛对策\n（项目）名称", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F1F5F9", is_label=True)
    table1.cell(2, 1).merge(table1.cell(2, 5))
    populate_cell(table1.cell(2, 1), "Rainbow-FinGPT：面向金融量化投研全流程的自主智能体系统", bold=True, font_size=9.2)

    # Row 3: 参赛组别
    populate_cell(table1.cell(3, 0), "参赛组别", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F1F5F9", is_label=True)
    table1.cell(3, 1).merge(table1.cell(3, 5))
    populate_cell(table1.cell(3, 1), "☑ 产教协同创新组    □ 区域特色产业组    □ 国产操作系统软件组", font_size=8.8)

    # Row 4-6: 项目负责人
    populate_cell(table1.cell(4, 0), "项目团队\n负责人及\n联系方式", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F1F5F9", is_label=True)
    populate_cell(table1.cell(4, 1), "项目负责人", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F8FAFC", is_label=True)
    populate_cell(table1.cell(4, 2), "吴宇轩 (学号: 20253803068)", font_size=8.8)
    populate_cell(table1.cell(4, 3), "所在学院", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F8FAFC", is_label=True)
    table1.cell(4, 4).merge(table1.cell(4, 5))
    populate_cell(table1.cell(4, 4), "华南师范大学 阿伯丁数据科学与人工智能学院", font_size=8.8)

    populate_cell(table1.cell(5, 1), "学历层次", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F8FAFC", is_label=True)
    populate_cell(table1.cell(5, 2), "2025级 本科生", font_size=8.8)
    populate_cell(table1.cell(5, 3), "就读专业", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F8FAFC", is_label=True)
    table1.cell(5, 4).merge(table1.cell(5, 5))
    populate_cell(table1.cell(5, 4), "信息管理与信息系统 (中外联合办学)", font_size=8.8)

    populate_cell(table1.cell(6, 1), "联系方式", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F8FAFC", is_label=True)
    table1.cell(6, 2).merge(table1.cell(6, 5))
    populate_cell(table1.cell(6, 2), "手机: 19098047316  |  邮箱: 85871865@qq.com\n在线演示: https://yuxuanwucn.github.io/stock-dashboard/", font_size=8.5)
    table1.cell(4, 0).merge(table1.cell(6, 0))

    # Row 7-9: 指导教师
    populate_cell(table1.cell(7, 0), "指导教师\n（5人以内）", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F1F5F9", is_label=True)
    populate_cell(table1.cell(7, 1), "姓名", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F8FAFC", is_label=True)
    table1.cell(7, 2).merge(table1.cell(7, 3))
    populate_cell(table1.cell(7, 2), "所在学院 / 单位", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F8FAFC", is_label=True)
    table1.cell(7, 4).merge(table1.cell(7, 5))
    populate_cell(table1.cell(7, 4), "职务 / 职称 / 联系方式", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F8FAFC", is_label=True)

    populate_cell(table1.cell(8, 1), "连洪泉 老师", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=8.8)
    table1.cell(8, 2).merge(table1.cell(8, 3))
    populate_cell(table1.cell(8, 2), "华南师范大学 经济与管理学院 / 科技商学院", font_size=8.6)
    table1.cell(8, 4).merge(table1.cell(8, 5))
    populate_cell(table1.cell(8, 4), "副教授 / 科技金融与经济学\n电话: 15626278766", font_size=8.6)

    populate_cell(table1.cell(9, 1), "企业导师组", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=8.8)
    table1.cell(9, 2).merge(table1.cell(9, 3))
    populate_cell(table1.cell(9, 2), "达观数据有限公司 (Datagrand Inc.)", font_size=8.6)
    table1.cell(9, 4).merge(table1.cell(9, 5))
    populate_cell(table1.cell(9, 4), "技术专家 / 金融大模型研发总监", font_size=8.6)
    table1.cell(7, 0).merge(table1.cell(9, 0))

    # Row 10-14: 项目团队全员信息
    populate_cell(table1.cell(10, 0), "项目团队\n全员信息\n(3-15人)\n【跨校协同】", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F1F5F9", is_label=True)
    populate_cell(table1.cell(10, 1), "成员姓名", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F8FAFC", is_label=True)
    populate_cell(table1.cell(10, 2), "团队职务与核心分工", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F8FAFC", is_label=True)
    populate_cell(table1.cell(10, 3), "所在院所及学历", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F8FAFC", is_label=True)
    populate_cell(table1.cell(10, 4), "年级及专业", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F8FAFC", is_label=True)
    populate_cell(table1.cell(10, 5), "学号", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F8FAFC", is_label=True)

    # Member 1: 吴宇轩
    populate_cell(table1.cell(11, 1), "吴宇轩", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=8.8)
    populate_cell(table1.cell(11, 2), "项目负责人 (系统总体架构 / 资产定价与风控总监 / 存储实证)", font_size=8.4)
    populate_cell(table1.cell(11, 3), "华南师大阿伯丁学院 本科", font_size=8.4)
    populate_cell(table1.cell(11, 4), "2025级信管", font_size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER)
    populate_cell(table1.cell(11, 5), "20253803068", font_size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Member 2: 阮奕霖
    populate_cell(table1.cell(12, 1), "阮奕霖", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=8.8)
    populate_cell(table1.cell(12, 2), "算法研发总监 (跨校技术负责人 / 多Agent博弈与状态机校准)", font_size=8.4)
    populate_cell(table1.cell(12, 3), "中山大学智工学院 本科", font_size=8.4)
    populate_cell(table1.cell(12, 4), "2025级智科", font_size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER)
    populate_cell(table1.cell(12, 5), "25361136", font_size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Member 3: 孟格漫
    populate_cell(table1.cell(13, 1), "孟格漫", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=8.8)
    populate_cell(table1.cell(13, 2), "数据工程总监 (知识抽取 / SCNU-RAG 事实锚点 / 研报解析)", font_size=8.4)
    populate_cell(table1.cell(13, 3), "华南师大阿伯丁学院 本科", font_size=8.4)
    populate_cell(table1.cell(13, 4), "2025级信管", font_size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER)
    populate_cell(table1.cell(13, 5), "20253803070", font_size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Member 4: 江昊
    populate_cell(table1.cell(14, 1), "江昊", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=8.8)
    populate_cell(table1.cell(14, 2), "量化回测总监 (风控工程 / 黄金与绿电跨板块实证检验)", font_size=8.4)
    populate_cell(table1.cell(14, 3), "华南师大阿伯丁学院 本科", font_size=8.4)
    populate_cell(table1.cell(14, 4), "2025级信管", font_size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER)
    populate_cell(table1.cell(14, 5), "20253803022", font_size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER)

    table1.cell(10, 0).merge(table1.cell(14, 0))

    # Column widths for Table 1: Total = 7.00 inches
    # Col 0: 1.00 in, Col 1: 0.80 in, Col 2: 2.10 in, Col 3: 1.25 in, Col 4: 0.85 in, Col 5: 1.00 in
    widths_t1 = [Inches(1.00), Inches(0.80), Inches(2.10), Inches(1.25), Inches(0.85), Inches(1.00)]
    for row in table1.rows:
        for idx, w in enumerate(widths_t1):
            if idx < len(row.cells):
                row.cells[idx].width = w

    # -------------------------------------------------------------
    # PAGE BREAK: Clean cut to Page 2
    # -------------------------------------------------------------
    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 2: Table 2 (Analysis, Strategy, Evidence)
    # -------------------------------------------------------------
    table2 = doc.add_table(rows=3, cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.autofit = False
    set_table_borders(table2)

    for row in table2.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

    for row in table2.rows:
        row.cells[0].width = Inches(1.25)
        row.cells[1].width = Inches(5.75)

    # Row 0: 命题分析与项目介绍 (300字以内)
    populate_cell(table2.cell(0, 0), "命题分析\n（300字以内）", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F1F5F9", is_label=True, pad_top=70, pad_bottom=70)
    p_intro = (
        "针对金融量化投研中初级研究员人力投入大、研报复现周期长（4-20h/篇）、通用大模型存在严重数值幻觉与前视未来函数泄漏等行业痛点，"
        "本项目针对达观数据产业命题，首创【FinRobot多专家审议 -> FinGPT领域后训练 -> NALE资源图谱 -> KHunter计量模拟器】四位一体全流程技术链。\n"
        "系统在存储超级周期、黄金事件驱动与绿电公用事业三大极端板块中完成样本外物理隔离拟真交易实测，全方位击败对应金融行业 ETF，回撤实现腰斩压制，调仓摩擦仅 0.15%（免除公募 1.5%~2% 管理费）。"
        "实测研报提取准确率 92.4%，代码运行率 98.9%，端到端投研耗时缩短 85% 以上，全面超额达成达观数据 10 项考核指标。"
    )
    populate_cell(table2.cell(0, 1), p_intro, font_size=9.0, line_spacing=1.20, pad_top=70, pad_bottom=70)

    # Row 1: 对策简介 (200字以内)
    populate_cell(table2.cell(1, 0), "对策简介\n（200字以内）", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F1F5F9", is_label=True, pad_top=70, pad_bottom=70)
    p_strategy = (
        "本项目采用四位一体自主决策闭环解决方案：\n"
        "① FinRobot 多专家博弈审议：宏观+行业+风控多角色动态辩论，剔除情绪偏误；\n"
        "② FinGPT 领域后训练：结合 SCNU-RAG 抽取带坐标事实三元组，拒绝臆造；\n"
        "③ NALE 产业与资源网络拓扑传导：领先卖方研报 5 日捕捉上游调价，剥离稳健特质 Alpha；\n"
        "④ KHunter 纯数学计量引擎：采用 Fama-MacBeth 3.0 回归 + HAC 稳健性检验 + Trend Gate 因果状态机，极端暴跌与 C 浪破位强制清仓，兼具高弹性与硬核防守。"
    )
    populate_cell(table2.cell(1, 1), p_strategy, font_size=9.0, line_spacing=1.20, pad_top=70, pad_bottom=70)

    # Row 2: 佐证材料说明
    populate_cell(table2.cell(2, 0), "佐证材料说明\n（知识产权等）", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color="F1F5F9", is_label=True, pad_top=70, pad_bottom=70)
    p_evidence = (
        "1. 出版级实证研报成果库：3 篇 3 页出版级高清实证研报 PDF（存储超级周期、黄金地缘避险、绿电公用事业，集成微观财务勾稽矩阵、波浪防御与 HAC 计量显著性检验）；\n"
        "2. 实证方法论学术白皮书：《Rainbow-FinGPT 数据溯源、标的池设计与实证方法论学术白皮书》，阐明多源数据采集流、双层证据金字塔与 Wind/CSMAR 迁移契约；\n"
        "3. 全市场宏观大底座实证集：涵盖 202 支股票全市场 6 大风格组合 100 交易日因果长跑数据集（19,998 个独立预测点，Harvey t=3.85，Brier=0.2481）；\n"
        "4. 软件著作权与代码工程：Rainbow-FinGPT v2.0 智能体中枢系统源码，包含 90+ 项全量自动化 pytest 单元与集成测试套件及无人值守自动化跑批；\n"
        "5. 达观数据产学研协同证明：针对达观曹植大模型金融垂直量化投研插件的技术接口协议与联合实施方案。"
    )
    populate_cell(table2.cell(2, 1), p_evidence, font_size=8.6, line_spacing=1.18, pad_top=70, pad_bottom=70)

    return doc

if __name__ == "__main__":
    print("Script build_perfect_attachment3 ready")
